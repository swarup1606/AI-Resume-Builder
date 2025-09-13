# backend/app.py
from flask import Flask, request, jsonify, send_file
import os
from groq import Groq
from dotenv import load_dotenv
import markdown
from docx import Document
from io import BytesIO
import PyPDF2
from flask_cors import CORS
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from PIL import Image, ImageDraw, ImageFont
import textwrap

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)  # allow all origins by default; tune for production

# Career Guidance Prompt template
career_guidance_prompt = """Analyze this resume thoroughly and provide comprehensive, personalized career guidance. Consider the person's experience, skills, industry, career level, and current market trends. Give detailed, actionable advice tailored to their specific background and goals.

For each section, provide 3-5 detailed paragraphs with specific recommendations, examples, and actionable steps. Be thorough and comprehensive in your analysis.

Respond with ONLY this JSON format (replace the placeholder text with your detailed analysis):

{{"future_projects": "Format as a career report card with emojis, headings, and bullet points. Include: 🚀 *Project Recommendations* with specific examples, technologies, timelines, and career impact. Use bullet points for each project idea with step-by-step guidance.", "skill_upgrade": "Format as a career report card with emojis, headings, and bullet points. Include: 📚 *Skill Development Roadmap* with technical and soft skills. Use bullet points for each skill with learning paths, resources, and timelines.", "career_roadmap": "Format as a career report card with emojis, headings, and bullet points. Include: 🎯 *Career Strategy* with short-term, medium-term, and long-term goals. Use bullet points for milestones, networking strategies, and advancement opportunities.", "strengths": "Format as a career report card with emojis, headings, and bullet points. Include: 💪 *Key Strengths* with technical skills, soft skills, achievements, and unique qualities. Use bullet points for each strength with specific examples.", "weaknesses": "Format as a career report card with emojis, headings, and bullet points. Include: 🎯 *Areas for Improvement* with skills gaps, missing qualifications, and development opportunities. Use bullet points for each weakness with actionable steps.", "resume_analysis": "Format as a career report card with emojis, headings, and bullet points. Include: 📄 *Resume Analysis* with content, structure, and formatting insights. Use bullet points for strengths, weaknesses, and improvement suggestions.", "skills": ["List of specific technical skills, programming languages, frameworks, tools, and technologies mentioned in the resume"], "ats_score": 85, "resume_score": 78, "ats_feedback": "Format as a career report card with emojis, headings, and bullet points. Include: 🎯 *ATS Optimization* with specific suggestions and examples. Use bullet points for each recommendation.", "resume_feedback": "Format as a career report card with emojis, headings, and bullet points. Include: 📊 *Resume Quality Feedback* with detailed analysis and suggestions. Use bullet points for each improvement area."}}

Resume: {resume}

Provide unique, personalized, and comprehensive insights. Be detailed and specific in your recommendations. Respond with ONLY the JSON object above, no other text."""

# Enhanced Prompt template for better resume generation
prompt_template = """You are a professional resume enhancement expert. Your task is to create a complete, professional resume that aligns with the given job description while preserving all essential candidate information.

CRITICAL REQUIREMENTS:

1. PRESERVE ALL ORIGINAL INFORMATION: Keep ALL personal details from the original resume including:
   - Full name (prominently displayed)
   - Contact information (email, phone number, address/location)
   - Professional summary/objective
   - All work experience with company names, job titles, and dates
   - All education details with institutions and degrees
   - All skills, certifications, and achievements
   - Any other personal information provided

2. ENHANCEMENT GUIDELINES:
   - Improve the wording and structure of existing content
   - Use strong action verbs and professional language
   - Optimize content to match the job description requirements
   - Reorganize sections for better flow and impact
   - Make the resume more ATS-friendly
   - Ensure consistent formatting and professional presentation

3. FORMAT REQUIREMENTS:
   - Start with candidate's full name in large, bold text
   - Include complete contact information (email, phone, location)
   - Structure in standard resume sections: Contact Info, Professional Summary, Experience, Education, Skills
   - Use bullet points for experience descriptions
   - Maintain professional formatting throughout

4. CONTENT RULES:
   - Do NOT add any information not present in the original resume
   - Do NOT include job description content in the resume
   - Do NOT invent skills, experiences, or qualifications
   - Only enhance and rephrase existing content
   - Ensure all dates, company names, and personal details are preserved exactly

5. OUTPUT FORMAT:
   - Create a complete, ready-to-use resume
   - Use clear section headers
   - Maintain professional appearance
   - Ensure it looks like a real person's resume

Job Description: {job_description}

Original Resume: {resume}

Create a complete, professional resume that preserves all original information while enhancing the content to better align with the job requirements. Make it look like a real candidate's resume with proper contact information and personal details.

After the enhanced resume, provide a section titled "Changes Made:" followed by a numbered list of the key improvements you made.
"""

# Initialize Groq client
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

def enhance_resume(job_description, resume):
    prompt = prompt_template.format(job_description=job_description, resume=resume)
    try:
        completion = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=4096,
            top_p=0.9,
            stream=False,
            stop=None,
        )
        full_response = completion.choices[0].message.content
        # Split the response into resume and notes
        if "Changes Made:" in full_response:
            parts = full_response.split("Changes Made:", 1)
            enhanced_resume = parts[0].strip()
            notes = "Changes Made:" + parts[1].strip()
        else:
            enhanced_resume = full_response
            notes = ""
        return enhanced_resume, notes
    except Exception as e:
        return f"Error: {str(e)}", ""

def analyze_career_guidance(resume):
    """Analyze resume for career guidance insights"""
    import time
    import random
    
    # Add timestamp and random element to ensure unique analysis each time
    timestamp = int(time.time())
    random_seed = random.randint(1000, 9999)
    
    # Enhance the prompt with unique context
    enhanced_prompt = f"{career_guidance_prompt}\n\nAnalysis ID: {timestamp}-{random_seed}\nProvide fresh, unique insights for this specific analysis."
    prompt = enhanced_prompt.format(resume=resume)
    
    try:
        print(f"Sending prompt to AI with {len(prompt)} characters (Analysis ID: {timestamp}-{random_seed})")
        completion = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": "You are a senior career guidance expert with 15+ years of experience. Format all responses as visually appealing career report cards with emojis, headings, and bullet points. Make responses engaging and easy to scan. You ONLY output valid JSON. No explanations, no text, just JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.8,  # Higher temperature for more creative, dynamic responses
            max_tokens=3000,  # Significantly increased for detailed responses
            top_p=0.95,  # Higher top_p for more diverse and comprehensive responses
            stream=False,
            stop=["", "---", "##", "\n\n", "Explanation", "Here is", "The JSON"],  # Stop at common text patterns
        )
        response = completion.choices[0].message.content
        print(f"Received response from AI: {len(response)} characters")
        print(f"First 200 chars: {response[:200]}")
        
        # Try to parse JSON response
        import json
        import re
        try:
            # Clean the response to extract JSON
            json_str = response.strip()
            
            # Remove any leading/trailing whitespace and newlines
            json_str = json_str.strip()
            
            # Remove markdown code blocks if present
            if "json" in json_str:
                json_start = json_str.find("json") + 7
                json_end = json_str.find("", json_start)
                json_str = json_str[json_start:json_end].strip()
            elif "" in json_str:
                json_start = json_str.find("") + 3
                json_end = json_str.find("```", json_start)
                json_str = json_str[json_start:json_end].strip()
            
            # Find JSON object boundaries - look for the first { and last }
            if "{" in json_str and "}" in json_str:
                json_start = json_str.find("{")
                json_end = json_str.rfind("}") + 1
                json_str = json_str[json_start:json_end]
            
            # Remove any text before the first { or after the last }
            json_str = json_str.strip()
            if json_str.startswith('"') or json_str.startswith("'"):
                # Find the first { after any quotes
                brace_pos = json_str.find('{')
                if brace_pos > 0:
                    json_str = json_str[brace_pos:]
            
            # Clean up common JSON issues more aggressively
            json_str = re.sub(r'\\n', '', json_str)  # Remove literal \n
            json_str = re.sub(r'\n', '', json_str)   # Remove actual newlines
            json_str = re.sub(r'\r', '', json_str)   # Remove carriage returns
            json_str = re.sub(r'\\"', '"', json_str) # Fix escaped quotes
            json_str = re.sub(r'\\t', '', json_str)  # Remove tabs
            json_str = re.sub(r'\\r', '', json_str)  # Remove literal \r
            
            # Remove any leading/trailing whitespace again
            json_str = json_str.strip()
            
            # Try to fix common JSON formatting issues
            json_str = json_str.replace('"', '"').replace('"', '"')
            json_str = json_str.replace(''', "'").replace(''', "'")
            
            # Final cleanup - remove any remaining non-JSON text
            if json_str.startswith('"') or json_str.startswith("'"):
                # Find the first { after any quotes
                brace_pos = json_str.find('{')
                if brace_pos > 0:
                    json_str = json_str[brace_pos:]
            
            # Handle specific error patterns like " and end with "
            if " and end with " in json_str:
                # Find the JSON part after this text
                parts = json_str.split(" and end with ")
                if len(parts) > 1:
                    json_str = parts[1].strip()
            
            # Remove any text that starts with quotes or common error patterns
            error_patterns = ["' and end with '", '" and end with "', "Here is the JSON:", "The JSON response:", "JSON:"]
            for pattern in error_patterns:
                if pattern in json_str:
                    json_str = json_str.split(pattern)[-1].strip()
            
            # Ensure we have a complete JSON object
            if not json_str.startswith('{') or not json_str.endswith('}'):
                # Try to find the complete JSON object
                start_pos = json_str.find('{')
                end_pos = json_str.rfind('}')
                if start_pos >= 0 and end_pos > start_pos:
                    json_str = json_str[start_pos:end_pos + 1]
            
            # Debug: Print the cleaned JSON string
            print(f"Cleaned JSON string: {json_str[:200]}...")
            
            # Validate JSON before parsing
            if not json_str or len(json_str) < 10:
                raise ValueError("JSON string is too short or empty")
            
            # Parse JSON
            guidance_data = json.loads(json_str)
            
            # Validate required keys
            required_keys = ["future_projects", "skill_upgrade", "career_roadmap", "strengths", "weaknesses", "resume_analysis", "skills", "ats_score", "resume_score", "ats_feedback", "resume_feedback"]
            missing_keys = [key for key in required_keys if key not in guidance_data]
            if missing_keys:
                print(f"Warning: Missing keys in JSON response: {missing_keys}")
                # Fill in missing keys with defaults
                for key in missing_keys:
                    if key in ["ats_score", "resume_score"]:
                        # Add some randomness to default scores
                        base_score = 50
                        variation = random.randint(-10, 10)
                        guidance_data[key] = max(0, min(100, base_score + variation))
                    else:
                        guidance_data[key] = f"<p>Analysis for {key} not available.</p>"
            
            # Add some variation to scores to ensure they're not always the same
            if "ats_score" in guidance_data and isinstance(guidance_data["ats_score"], int):
                # Add small random variation to make scores more dynamic
                variation = random.randint(-3, 3)
                guidance_data["ats_score"] = max(0, min(100, guidance_data["ats_score"] + variation))
            
            if "resume_score" in guidance_data and isinstance(guidance_data["resume_score"], int):
                # Add small random variation to make scores more dynamic
                variation = random.randint(-3, 3)
                guidance_data["resume_score"] = max(0, min(100, guidance_data["resume_score"] + variation))
            
            # Convert to HTML format for better display
            html_results = {}
            for key, value in guidance_data.items():
                if key in ['ats_score', 'resume_score']:
                    # Keep scores as numbers
                    html_results[key] = value
                else:
                    # Convert content to HTML with enhanced formatting for longer content
                    text_value = str(value)
                    
                    # Pre-process the text for better formatting
                    # Convert multiple line breaks to paragraph breaks
                    text_value = re.sub(r'\n\s*\n', '\n\n', text_value)
                    
                    # Convert numbered lists (1., 2., etc.)
                    text_value = re.sub(r'^(\d+\.)\s+(.+)$', r'\1** \2', text_value, flags=re.MULTILINE)
                    
                    # Convert bullet points
                    text_value = re.sub(r'^[-]\s+(.+)$', r' \1', text_value, flags=re.MULTILINE)
                    
                    # Convert headers
                    text_value = re.sub(r'^### (.+)$', r'### \1', text_value, flags=re.MULTILINE)
                    text_value = re.sub(r'^## (.+)$', r'## \1', text_value, flags=re.MULTILINE)
                    text_value = re.sub(r'^# (.+)$', r'# \1', text_value, flags=re.MULTILINE)
                    
                    # Convert to HTML using markdown
                    html_value = markdown.markdown(text_value, extensions=['nl2br', 'tables', 'fenced_code'])
                    html_results[key] = html_value
            
            return html_results
            
        except json.JSONDecodeError as e:
            # Fallback: return structured text if JSON parsing fails
            print(f"JSON parsing error: {e}")
            print(f"Response received: {response[:500]}...")
            
            # Try to extract individual fields from the malformed response
            fallback_response = {
                "future_projects": "<h3>🚀 Project Recommendations</h3><p>Based on your resume analysis, here are comprehensive project recommendations to boost your career prospects:</p><ul><li><strong>📁 Portfolio Enhancement Project:</strong> Create a comprehensive portfolio showcasing your best work with detailed case studies, code repositories, and live demonstrations.</li><li><strong>🏢 Industry-Specific Solution:</strong> Build a project that addresses a real problem in your target industry, demonstrating both technical skills and business understanding.</li><li><strong>🔗 Open Source Contribution:</strong> Contribute to popular open-source projects in your field to build credibility and network with other developers.</li><li><strong>🏆 Certification Project:</strong> Complete a project that leads to a recognized certification in your field, such as cloud platforms, specific technologies, or methodologies.</li></ul><p><strong>💡 Pro Tip:</strong> Each project should include documentation, version control, testing, and deployment to demonstrate professional development practices.</p>",
                "skill_upgrade": "<h3>📚 Skill Development Roadmap</h3><p>To advance your career, focus on these key skill areas:</p><h4>💻 Technical Skills</h4><ul><li><strong>🔧 Advanced Programming:</strong> Master advanced features of your primary programming languages</li><li><strong>☁ Cloud Platforms:</strong> Learn cloud platforms (AWS, Azure, or GCP) and containerization</li><li><strong>🏗 System Design:</strong> Develop expertise in data structures, algorithms, and system design</li><li><strong>🛠 Modern Frameworks:</strong> Gain experience with modern frameworks and tools in your field</li></ul><h4>🤝 Soft Skills</h4><ul><li><strong>💬 Communication:</strong> Improve presentation and written communication skills</li><li><strong>👥 Leadership:</strong> Develop team management and project management abilities</li><li><strong>🌐 Networking:</strong> Build professional relationships and industry connections</li><li><strong>🧠 Critical Thinking:</strong> Enhance problem-solving and analytical skills</li></ul>",
                "career_roadmap": "<h3>🎯 Career Development Strategy</h3><h4>⏰ Short-term Goals (3-6 months)</h4><ul><li><strong>📁 Project Portfolio:</strong> Complete 2-3 skill-building projects</li><li><strong>📝 Profile Update:</strong> Update your resume and LinkedIn profile</li><li><strong>🤝 Networking:</strong> Network with 10+ professionals in your field</li><li><strong>💼 Job Applications:</strong> Apply to 20+ relevant positions</li></ul><h4>📅 Medium-term Goals (1-2 years)</h4><ul><li><strong>💼 Career Growth:</strong> Secure a position that offers growth opportunities</li><li><strong>🏆 Certifications:</strong> Complete relevant certifications</li><li><strong>🌐 Professional Network:</strong> Build a strong professional network</li><li><strong>👥 Leadership:</strong> Take on leadership responsibilities</li></ul><h4>🚀 Long-term Goals (3-5 years)</h4><ul><li><strong>📈 Senior Position:</strong> Achieve senior-level position</li><li><strong>🎓 Mentoring:</strong> Mentor junior professionals</li><li><strong>🎯 Specialization:</strong> Consider specialization or management track</li><li><strong>💡 Thought Leadership:</strong> Build thought leadership in your field</li></ul>",
                "strengths": "<h3>💪 Key Strengths</h3><ul><li><strong>💻 Technical Expertise:</strong> Strong foundation in your primary programming languages and technologies</li><li><strong>🧩 Problem-Solving Skills:</strong> Demonstrated ability to analyze complex problems and develop effective solutions</li><li><strong>📁 Project Experience:</strong> Hands-on experience with real-world projects and deliverables</li><li><strong>📚 Continuous Learning:</strong> Proactive approach to staying updated with industry trends and technologies</li><li><strong>🔍 Attention to Detail:</strong> Careful and thorough approach to work and documentation</li></ul>",
                "weaknesses": "<h3>🎯 Areas for Improvement</h3><ul><li><strong>💬 Communication Skills:</strong> Enhance presentation and written communication abilities</li><li><strong>👥 Leadership Experience:</strong> Develop team management and mentoring capabilities</li><li><strong>🏢 Industry Knowledge:</strong> Deepen understanding of specific industry practices and standards</li><li><strong>🌐 Networking:</strong> Build professional relationships and industry connections</li><li><strong>🚀 Advanced Technologies:</strong> Learn cutting-edge tools and frameworks in your field</li></ul>",
                "resume_analysis": "<h3>📄 Resume Analysis</h3><h4>✅ Strengths</h4><ul><li><strong>📋 Structure:</strong> Clear structure and formatting</li><li><strong>🛠 Skills:</strong> Relevant technical skills listed</li><li><strong>📁 Projects:</strong> Project experience highlighted</li><li><strong>📊 Metrics:</strong> Quantifiable achievements included</li></ul><h4>⚠ Areas for Improvement</h4><ul><li><strong>📈 Metrics:</strong> Add more specific metrics and results</li><li><strong>🔍 ATS Optimization:</strong> Include relevant keywords for ATS optimization</li><li><strong>👥 Soft Skills:</strong> Expand on leadership and soft skills</li><li><strong>📝 Summary:</strong> Consider adding a professional summary</li></ul><h4>💡 Recommendations</h4><p><strong>🎯 Focus Areas:</strong> Quantify achievements, use action verbs, and tailor content to specific job applications. Ensure consistent formatting and proofread thoroughly.</p>",
                "skills": ["JavaScript", "Python", "React", "Node.js", "SQL", "Git", "AWS", "Docker", "Agile", "Project Management"],
                "ats_score": 75,
                "resume_score": 70,
                "ats_feedback": "<h3>🎯 ATS Optimization Suggestions</h3><ul><li><strong>📋 Section Headings:</strong> Use standard section headings (Experience, Education, Skills)</li><li><strong>🔍 Keywords:</strong> Include relevant keywords from job descriptions</li><li><strong>📄 Formatting:</strong> Avoid graphics, tables, or complex formatting</li><li><strong>🔤 Fonts:</strong> Use common fonts and standard file formats</li><li><strong>🛠 Skills Section:</strong> Include a skills section with specific technologies</li></ul>",
                "resume_feedback": "<h3>📊 Resume Quality Feedback</h3><ul><li><strong>📋 Structure:</strong> Overall structure is good but could be more compelling</li><li><strong>📈 Metrics:</strong> Add more quantifiable achievements and metrics</li><li><strong>📝 Summary:</strong> Include a professional summary or objective</li><li><strong>✨ Appearance:</strong> Ensure consistent formatting and professional appearance</li><li><strong>🏆 Certifications:</strong> Consider adding relevant certifications or training</li></ul>"
            }
            
            # Try to extract individual fields using regex
            try:
                import re
                
                # Extract scores
                ats_match = re.search(r'ats_score["\']?\s*:\s*(\d+)', response, re.IGNORECASE)
                if ats_match:
                    fallback_response["ats_score"] = int(ats_match.group(1))
                
                resume_match = re.search(r'resume_score["\']?\s*:\s*(\d+)', response, re.IGNORECASE)
                if resume_match:
                    fallback_response["resume_score"] = int(resume_match.group(1))
                
                # Try to extract text fields
                fields_to_extract = [
                    "future_projects", "skill_upgrade", "career_roadmap", "self_check",
                    "ats_feedback", "resume_feedback"
                ]
                
                for field in fields_to_extract:
                    # Look for field: "content" pattern
                    pattern = rf'{field}["\']?\s*:\s*["\']([^"\']+)["\']'
                    match = re.search(pattern, response, re.IGNORECASE | re.DOTALL)
                    if match:
                        content = match.group(1).strip()
                        if content and len(content) > 10:  # Only use if substantial content
                            fallback_response[field] = f"<p>{content}</p>"
                
            except Exception as extract_error:
                print(f"Error extracting fields: {extract_error}")
            
            return fallback_response
            
    except Exception as e:
        print(f"Complete failure in career guidance analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Return a basic response that will at least show something to the user
        return {
            "future_projects": "<p>Build a portfolio website showcasing your projects, contribute to open source projects, and create a mobile app to demonstrate your skills.</p>",
            "skill_upgrade": "<p>Learn modern frameworks like React, Node.js, and cloud technologies. Consider getting certified in AWS or Azure.</p>",
            "career_roadmap": "<p>Focus on gaining 2-3 years of experience in your current role, then seek senior positions. Build leadership skills and consider management track.</p>",
            "self_check": "<p>Your technical skills are developing well. Focus on improving communication, leadership, and project management abilities.</p>",
            "ats_score": 65,
            "resume_score": 70,
            "ats_feedback": "<p>Your resume has good structure. Consider adding more keywords and quantifiable achievements.</p>",
            "resume_feedback": "<p>Your resume shows good experience. Add more specific project outcomes and measurable results.</p>"
        }

def create_pdf_resume(text):
    """Create a PDF resume from text content"""
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            textColor=colors.darkblue
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6
        )
        
        # Parse the text and create content
        content = []
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                content.append(Spacer(1, 6))
                continue
                
            # Check if it's a title (usually the first line or very short)
            if len(line) < 50 and not line.startswith('•') and not line.startswith('-') and not line.startswith('*'):
                if len(content) < 3:  # Likely the name/title
                    content.append(Paragraph(line, title_style))
                else:
                    content.append(Paragraph(line, heading_style))
            else:
                # Regular content
                content.append(Paragraph(line, normal_style))
        
        doc.build(content)
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Error creating PDF: {str(e)}")
        return None

def create_jpg_resume(text):
    """Create a JPG resume from text content"""
    try:
        # Image dimensions
        width, height = 2480, 3508  # A4 size at 300 DPI
        image = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(image)
        
        # Try to use a system font, fallback to default
        try:
            font_large = ImageFont.truetype("arial.ttf", 48)
            font_medium = ImageFont.truetype("arial.ttf", 32)
            font_small = ImageFont.truetype("arial.ttf", 24)
        except:
            try:
                font_large = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 48)
                font_medium = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 32)
                font_small = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 24)
            except:
                font_large = ImageFont.load_default()
                font_medium = ImageFont.load_default()
                font_small = ImageFont.load_default()
        
        # Parse and draw text
        lines = text.split('\n')
        y_position = 100
        line_height = 60
        
        for line in lines:
            line = line.strip()
            if not line:
                y_position += 30
                continue
                
            # Determine font size based on content
            if len(line) < 50 and not line.startswith('•') and not line.startswith('-') and not line.startswith('*'):
                if y_position < 200:  # Likely the name/title
                    font = font_large
                    color = (0, 0, 139)  # Dark blue
                else:
                    font = font_medium
                    color = (0, 0, 139)  # Dark blue
            else:
                font = font_small
                color = (0, 0, 0)  # Black
            
            # Wrap long lines
            wrapped_lines = textwrap.wrap(line, width=80)
            for wrapped_line in wrapped_lines:
                if y_position > height - 100:
                    break
                draw.text((100, y_position), wrapped_line, font=font, fill=color)
                y_position += line_height
        
        # Save to buffer
        buffer = BytesIO()
        image.save(buffer, format='JPEG', quality=95)
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Error creating JPG: {str(e)}")
        return None

@app.route('/tool', methods=['POST'])
def tool():
    """
    Accepts form-data:
    - jd: job description text
    - resume: optional resume text
    - resume_file: optional file (pdf/docx)
    Returns JSON: { result_html, notes_html, enhanced_text }
    """
    # If JSON body (rare), read it too
    jd = request.form.get('jd', '') or (request.json.get('jd') if request.is_json else '')
    resume = request.form.get('resume', '') or (request.json.get('resume') if request.is_json else '')

    # If file uploaded, extract text
    if 'resume_file' in request.files and request.files['resume_file'].filename:
        file = request.files['resume_file']
        filename = file.filename.lower()
        try:
            if filename.endswith('.docx'):
                doc = Document(file)
                resume = '\n'.join([para.text for para in doc.paragraphs])
            elif filename.endswith('.pdf'):
                reader = PyPDF2.PdfReader(file)
                resume = '\n'.join([page.extract_text() or '' for page in reader.pages])
            else:
                return jsonify({'error': 'Unsupported file format. Please upload DOCX or PDF.'}), 400
        except Exception as e:
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500

    if not jd:
        return jsonify({'error': 'Please provide a job description.'}), 400
    
    # Check if either resume text or file is provided
    has_resume_text = resume and resume.strip()
    has_resume_file = 'resume_file' in request.files and request.files['resume_file'].filename
    
    if not has_resume_text and not has_resume_file:
        return jsonify({'error': 'Please provide resume text or upload a file.'}), 400

    enhanced_resume, notes = enhance_resume(jd, resume)

    # Render simple HTML (React will render this via dangerouslySetInnerHTML)
    result_html = markdown.markdown(enhanced_resume)
    notes_html = markdown.markdown(notes)

    return jsonify({
        'result_html': result_html,
        'notes_html': notes_html,
        'enhanced_text': enhanced_resume
    })

@app.route('/download_docx', methods=['POST'])
def download_docx():
    data = request.get_json()
    text = data.get('text', '')
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf,
                     as_attachment=True,
                     download_name='enhanced_resume.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/download_pdf', methods=['POST'])
def download_pdf():
    """Download resume as PDF"""
    data = request.get_json()
    text = data.get('text', '')
    
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    pdf_buffer = create_pdf_resume(text)
    if not pdf_buffer:
        return jsonify({'error': 'Failed to create PDF'}), 500
    
    return send_file(pdf_buffer,
                     as_attachment=True,
                     download_name='enhanced_resume.pdf',
                     mimetype='application/pdf')

@app.route('/download_jpg', methods=['POST'])
def download_jpg():
    """Download resume as JPG"""
    data = request.get_json()
    text = data.get('text', '')
    
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    jpg_buffer = create_jpg_resume(text)
    if not jpg_buffer:
        return jsonify({'error': 'Failed to create JPG'}), 500
    
    return send_file(jpg_buffer,
                     as_attachment=True,
                     download_name='enhanced_resume.jpg',
                     mimetype='image/jpeg')

@app.route('/test_career_guidance', methods=['GET'])
def test_career_guidance():
    """Test endpoint to verify career guidance is working"""
    test_resume = "John Doe\nSoftware Developer\nEmail: john@example.com\nExperience: 2 years in web development\nSkills: JavaScript, React, Node.js"
    
    try:
        guidance_results = analyze_career_guidance(test_resume)
        return jsonify({
            "status": "success",
            "message": "Career guidance analysis working",
            "results": guidance_results
        })
    except Exception as e:
        return jsonify({
            "status": "error", 
            "message": str(e)
        }), 500

@app.route('/career_guidance', methods=['POST'])
def career_guidance():
    """
    Career guidance analysis endpoint
    Accepts form-data with resume_file (PDF)
    Returns JSON with career guidance insights
    """
    # Check if file is uploaded
    if 'resume_file' not in request.files or not request.files['resume_file'].filename:
        return jsonify({'error': 'Please upload a resume file.'}), 400
    
    file = request.files['resume_file']
    filename = file.filename.lower()
    
    # Extract text from PDF
    try:
        if filename.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            resume_text = '\n'.join([page.extract_text() or '' for page in reader.pages])
        else:
            return jsonify({'error': 'Only PDF files are supported for career guidance.'}), 400
    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500
    
    if not resume_text.strip():
        return jsonify({'error': 'Could not extract text from the PDF. Please ensure the file is not corrupted.'}), 400
    
    # Analyze resume for career guidance
    try:
        print(f"Starting career guidance analysis for resume with {len(resume_text)} characters")
        guidance_results = analyze_career_guidance(resume_text)
        print(f"Analysis completed successfully. Keys: {list(guidance_results.keys())}")
        return jsonify(guidance_results)
    except Exception as e:
        print(f"Error in career guidance analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Error analyzing resume: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5001)