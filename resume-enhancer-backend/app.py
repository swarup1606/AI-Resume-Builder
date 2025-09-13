# backend/app.py
from flask import Flask, request, jsonify, send_file
import os
from groq import Groq
from dotenv import load_dotenv
import markdown
from docx import Document
from io import BytesIO
import PyPDF2
from flask_cors import CORS
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from PIL import Image, ImageDraw, ImageFont
import textwrap

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)  # allow all origins by default; tune for production

# Enhanced Prompt template for intelligent job-specific resume generation
prompt_template = """You are an expert AI resume enhancement specialist with deep knowledge of ATS systems and hiring practices. Your task is to transform a general resume into a highly targeted, job-specific resume that will pass ATS screening and impress hiring managers.

CRITICAL INSTRUCTIONS - FOLLOW THIS EXACT STRUCTURE:

1. ALWAYS EXTRACT & PRESERVE BASIC INFO:
   - Extract and keep: Name, Phone, Email, LinkedIn, GitHub from the original resume
   - These should ALWAYS be included regardless of job role
   - Format: Name (large, bold), then contact details below

2. AI-GENERATED PROFESSIONAL SUMMARY:
   - Analyze the job role and JD to create a compelling 3-4 line summary
   - Highlight the candidate's most relevant strengths for THIS specific role
   - Include key skills and experience that match the job requirements
   - Use action words and quantifiable achievements when possible

3. SMART EXPERIENCE SELECTION (MAX 2 EXPERIENCES):
   - Find the 1-2 experiences most relevant to the job role and JD
   - If only 1 relevant experience exists, add the next closest one
   - For each selected experience, enhance with AI-generated bullet points that:
     * Use keywords from the job description
     * Quantify achievements with numbers/metrics
     * Show impact and results
     * Align with the target role's requirements
   - If no relevant experience exists, select the 2 most impressive ones and enhance them

4. SMART PROJECT SELECTION (MAX 2 PROJECTS):
   - Find 1-2 projects most relevant to the job role and JD
   - If only 1 relevant project exists, add the next closest one
   - For each selected project, enhance with AI-generated descriptions that:
     * Use technologies mentioned in the job description
     * Show problem-solving and technical skills
     * Include quantifiable results or impact
     * Align with the target role's requirements
   - If no projects exist, generate 1-2 relevant projects based on experience and job requirements

5. EDUCATION SECTION:
   - Keep education exactly as it appears in the original resume
   - No AI generation or enhancement needed
   - Include: Degree, Institution, Year, GPA (if mentioned)

6. FILTERED SKILLS SECTION:
   - ONLY include skills that are relevant to the job role and JD
   - If no relevant skills exist, include the closest/most transferable ones
   - Group into: Technical Skills, Soft Skills, Tools & Technologies
   - Use exact keywords from the job description when possible

7. FILTERED CERTIFICATIONS SECTION:
   - ONLY include certifications relevant to the job role and JD
   - If no relevant certifications exist, include the closest/most valuable ones
   - Include: Certification Name, Issuing Organization, Date (if available)

8. INTERESTS SECTION:
   - Keep interests exactly as they appear in the original resume
   - No AI generation or enhancement needed

9. ATS OPTIMIZATION:
   - Use exact keywords from the job description throughout
   - Include industry-standard terminology
   - Ensure proper section headings and formatting
   - Use consistent date formats and bullet points

10. PROFESSIONAL FORMATTING:
    - Structure: Contact Info → Professional Summary → Experience → Projects → Education → Skills → Certifications → Interests
    - Use bullet points for experience and project descriptions
    - Include quantifiable achievements and metrics
    - Keep formatting clean, professional, and ATS-friendly

Job Description: {job_description}
Target Job Title: {job_title}

Original Resume: {resume}

Create a complete, professional, job-specific resume that will help the candidate stand out and get interviews. The resume should look like it was specifically crafted for this exact job opportunity, while being truthful and based on the candidate's actual background.

After the enhanced resume, provide a section titled "Enhancement Summary:" followed by a detailed explanation of how you made the resume job-specific and what improvements were made.
"""

# Initialize Groq client
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

def parse_resume_structure(resume_text):
    """Parse resume text to extract structured information for better AI processing"""
    parsing_prompt = f"""
    Parse this resume text and extract structured information in the following format:
    
    Resume Text: {resume_text}
    
    Please extract and return:
    
    **PERSONAL INFORMATION:**
    - Name: [Full Name]
    - Phone: [Phone Number]
    - Email: [Email Address]
    - LinkedIn: [LinkedIn URL if present]
    - GitHub: [GitHub URL if present]
    - Location: [City, State/Country if present]
    
    **EXPERIENCE:**
    - Job Title 1: [Title] | Company: [Company] | Duration: [Start Date - End Date] | Description: [Brief description]
    - Job Title 2: [Title] | Company: [Company] | Duration: [Start Date - End Date] | Description: [Brief description]
    - [Continue for all experiences]
    
    **PROJECTS:**
    - Project 1: [Project Name] | Technologies: [Tech stack] | Description: [Brief description]
    - Project 2: [Project Name] | Technologies: [Tech stack] | Description: [Brief description]
    - [Continue for all projects]
    
    **EDUCATION:**
    - Degree: [Degree] | Institution: [Institution] | Year: [Year] | GPA: [GPA if mentioned]
    
    **SKILLS:**
    - Technical: [List all technical skills]
    - Soft: [List all soft skills]
    - Tools: [List all tools and technologies]
    
    **CERTIFICATIONS:**
    - Certification 1: [Name] | Issuer: [Organization] | Date: [Date if mentioned]
    - Certification 2: [Name] | Issuer: [Organization] | Date: [Date if mentioned]
    - [Continue for all certifications]
    
    **INTERESTS:**
    - [List all interests mentioned]
    
    If any section is missing or unclear, mark it as "Not Found" or "Unclear".
    """
    
    try:
        completion = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": parsing_prompt}],
            temperature=0.1,
            max_tokens=2000,
            top_p=0.9,
            stream=False,
            stop=None,
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Parsing error: {str(e)}"

def analyze_job_requirements(job_description, job_title):
    """Analyze job description to extract key requirements and skills for intelligent resume enhancement"""
    analysis_prompt = f"""
    Analyze this job description to extract key information for intelligent resume enhancement:
    
    Job Title: {job_title}
    Job Description: {job_description}
    
    Please extract and return:
    1. **Required Technical Skills** - List all technical skills, programming languages, frameworks, tools mentioned
    2. **Required Soft Skills** - Leadership, communication, problem-solving, etc.
    3. **Key Technologies & Tools** - Specific software, platforms, methodologies mentioned
    4. **Important Keywords for ATS** - Exact phrases and terminology to include
    5. **Experience Level Required** - Entry, Mid, Senior level indicators
    6. **Industry/Company Type** - Tech, finance, healthcare, startup, enterprise, etc.
    7. **Key Responsibilities** - Main duties and expectations
    8. **Preferred Qualifications** - Nice-to-have skills and experiences
    9. **Project Types** - What kind of projects would be most relevant
    10. **Certification Preferences** - Any specific certifications mentioned
    
    Format as a structured analysis that will help the AI select the most relevant experiences, projects, skills, and certifications from the candidate's resume.
    """
    
    try:
        completion = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": analysis_prompt}],
            temperature=0.1,
            max_tokens=1500,
            top_p=0.9,
            stream=False,
            stop=None,
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Analysis error: {str(e)}"

def enhance_resume(job_description, resume, job_title=""):
    prompt = prompt_template.format(
        job_description=job_description, 
        resume=resume, 
        job_title=job_title
    )
    
    # Try with the more powerful model first
    try:
        completion = client.chat.completions.create(
            model="llama-3.1-70b-versatile",  # Using more powerful model for better results
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,  # Slightly higher for more creative but relevant content
            max_tokens=6144,  # Increased token limit for comprehensive resume generation
            top_p=0.9,
            stream=False,
            stop=None,
        )
        full_response = completion.choices[0].message.content
        # Split the response into resume and notes
        if "Enhancement Summary:" in full_response:
            parts = full_response.split("Enhancement Summary:", 1)
            enhanced_resume = parts[0].strip()
            notes = "Enhancement Summary:" + parts[1].strip()
        elif "Changes Made:" in full_response:
            parts = full_response.split("Changes Made:", 1)
            enhanced_resume = parts[0].strip()
            notes = "Changes Made:" + parts[1].strip()
        else:
            enhanced_resume = full_response
            notes = ""
        return enhanced_resume, notes
    except Exception as e:
        print(f"Error with 70b model: {str(e)}")
        # Fallback to the instant model if the 70b model fails
        try:
            completion = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=4096,
                top_p=0.9,
                stream=False,
                stop=None,
            )
            full_response = completion.choices[0].message.content
            # Split the response into resume and notes
            if "Enhancement Summary:" in full_response:
                parts = full_response.split("Enhancement Summary:", 1)
                enhanced_resume = parts[0].strip()
                notes = "Enhancement Summary:" + parts[1].strip()
            elif "Changes Made:" in full_response:
                parts = full_response.split("Changes Made:", 1)
                enhanced_resume = parts[0].strip()
                notes = "Changes Made:" + parts[1].strip()
            else:
                enhanced_resume = full_response
                notes = ""
            return enhanced_resume, notes
        except Exception as e2:
            print(f"Error with 8b model: {str(e2)}")
            return f"Error: Unable to process resume enhancement. Please try again. Error details: {str(e2)}", ""

def create_pdf_resume(text):
    """Create a PDF resume from text content"""
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            textColor=colors.darkblue
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6
        )
        
        # Parse the text and create content
        content = []
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                content.append(Spacer(1, 6))
                continue
                
            # Check if it's a title (usually the first line or very short)
            if len(line) < 50 and not line.startswith('•') and not line.startswith('-') and not line.startswith('*'):
                if len(content) < 3:  # Likely the name/title
                    content.append(Paragraph(line, title_style))
                else:
                    content.append(Paragraph(line, heading_style))
            else:
                # Regular content
                content.append(Paragraph(line, normal_style))
        
        doc.build(content)
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Error creating PDF: {str(e)}")
        return None

def create_jpg_resume(text):
    """Create a JPG resume from text content"""
    try:
        # Image dimensions
        width, height = 2480, 3508  # A4 size at 300 DPI
        image = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(image)
        
        # Try to use a system font, fallback to default
        try:
            font_large = ImageFont.truetype("arial.ttf", 48)
            font_medium = ImageFont.truetype("arial.ttf", 32)
            font_small = ImageFont.truetype("arial.ttf", 24)
        except:
            try:
                font_large = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 48)
                font_medium = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 32)
                font_small = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 24)
            except:
                font_large = ImageFont.load_default()
                font_medium = ImageFont.load_default()
                font_small = ImageFont.load_default()
        
        # Parse and draw text
        lines = text.split('\n')
        y_position = 100
        line_height = 60
        
        for line in lines:
            line = line.strip()
            if not line:
                y_position += 30
                continue
                
            # Determine font size based on content
            if len(line) < 50 and not line.startswith('•') and not line.startswith('-') and not line.startswith('*'):
                if y_position < 200:  # Likely the name/title
                    font = font_large
                    color = (0, 0, 139)  # Dark blue
                else:
                    font = font_medium
                    color = (0, 0, 139)  # Dark blue
            else:
                font = font_small
                color = (0, 0, 0)  # Black
            
            # Wrap long lines
            wrapped_lines = textwrap.wrap(line, width=80)
            for wrapped_line in wrapped_lines:
                if y_position > height - 100:
                    break
                draw.text((100, y_position), wrapped_line, font=font, fill=color)
                y_position += line_height
        
        # Save to buffer
        buffer = BytesIO()
        image.save(buffer, format='JPEG', quality=95)
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Error creating JPG: {str(e)}")
        return None

@app.route('/tool', methods=['POST'])
def tool():
    """
    Accepts form-data:
    - jd: job description text
    - resume: optional resume text
    - resume_file: optional file (pdf/docx)
    Returns JSON: { result_html, notes_html, enhanced_text }
    """
    # If JSON body (rare), read it too
    jd = request.form.get('jd', '') or (request.json.get('jd') if request.is_json else '')
    resume = request.form.get('resume', '') or (request.json.get('resume') if request.is_json else '')
    job_title = request.form.get('job_title', '') or (request.json.get('job_title') if request.is_json else '')

    # If file uploaded, extract text
    if 'resume_file' in request.files and request.files['resume_file'].filename:
        file = request.files['resume_file']
        filename = file.filename.lower()
        try:
            if filename.endswith('.docx'):
                doc = Document(file)
                resume = '\n'.join([para.text for para in doc.paragraphs])
            elif filename.endswith('.pdf'):
                reader = PyPDF2.PdfReader(file)
                resume = '\n'.join([page.extract_text() or '' for page in reader.pages])
            else:
                return jsonify({'error': 'Unsupported file format. Please upload DOCX or PDF.'}), 400
        except Exception as e:
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500

    if not jd:
        return jsonify({'error': 'Please provide a job description.'}), 400
    
    # Check if either resume text or file is provided
    has_resume_text = resume and resume.strip()
    has_resume_file = 'resume_file' in request.files and request.files['resume_file'].filename
    
    if not has_resume_text and not has_resume_file:
        return jsonify({'error': 'Please provide resume text or upload a file.'}), 400

    # Parse resume structure for better AI processing
    try:
        resume_structure = parse_resume_structure(resume)
    except Exception as e:
        print(f"Error parsing resume structure: {str(e)}")
        resume_structure = "Resume structure parsing failed. Proceeding with enhancement..."
    
    # Analyze job requirements for better targeting
    try:
        job_analysis = analyze_job_requirements(jd, job_title)
    except Exception as e:
        print(f"Error analyzing job requirements: {str(e)}")
        job_analysis = "Job analysis failed. Proceeding with enhancement..."
    
    # Enhance resume with job-specific intelligence
    enhanced_resume, notes = enhance_resume(jd, resume, job_title)

    # Render simple HTML (React will render this via dangerouslySetInnerHTML)
    result_html = markdown.markdown(enhanced_resume)
    notes_html = markdown.markdown(notes)
    
    # Include job analysis and resume structure in the response
    analysis_html = markdown.markdown(f"**Job Analysis:**\n{job_analysis}")
    structure_html = markdown.markdown(f"**Resume Structure Analysis:**\n{resume_structure}")

    return jsonify({
        'result_html': result_html,
        'notes_html': notes_html,
        'enhanced_text': enhanced_resume,
        'job_analysis': analysis_html,
        'resume_structure': structure_html
    })

@app.route('/download_docx', methods=['POST'])
def download_docx():
    data = request.get_json()
    text = data.get('text', '')
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf,
                     as_attachment=True,
                     download_name='enhanced_resume.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/download_pdf', methods=['POST'])
def download_pdf():
    """Download resume as PDF"""
    data = request.get_json()
    text = data.get('text', '')
    
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    pdf_buffer = create_pdf_resume(text)
    if not pdf_buffer:
        return jsonify({'error': 'Failed to create PDF'}), 500
    
    return send_file(pdf_buffer,
                     as_attachment=True,
                     download_name='enhanced_resume.pdf',
                     mimetype='application/pdf')

@app.route('/download_jpg', methods=['POST'])
def download_jpg():
    """Download resume as JPG"""
    data = request.get_json()
    text = data.get('text', '')
    
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    jpg_buffer = create_jpg_resume(text)
    if not jpg_buffer:
        return jsonify({'error': 'Failed to create JPG'}), 500
    
    return send_file(jpg_buffer,
                     as_attachment=True,
                     download_name='enhanced_resume.jpg',
                     mimetype='image/jpeg')

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)